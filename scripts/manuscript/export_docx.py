#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]


@dataclass(frozen=True)
class FigureCaption:
    figure_id: str  # e.g. "F1"
    title_line: str  # single-line caption header (without markdown **)
    body: str  # remaining caption body text


def _strip_md_inline(text: str) -> str:
    text = text.replace("\u2011", "-")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def _iter_blocks(md: str) -> list[tuple[str, list[str]]]:
    lines = md.splitlines()
    blocks: list[tuple[str, list[str]]] = []
    i = 0
    in_code = False
    code_buf: list[str] = []
    para_buf: list[str] = []

    def flush_para() -> None:
        nonlocal para_buf
        if para_buf:
            blocks.append(("para", para_buf))
            para_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                blocks.append(("code", code_buf))
                code_buf = []
                in_code = False
            else:
                flush_para()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line.rstrip("\n"))
            i += 1
            continue

        # skip leading YAML front matter if present
        if i == 0 and line.strip() == "---":
            j = 1
            while j < len(lines) and lines[j].strip() != "---":
                j += 1
            if j < len(lines) and lines[j].strip() == "---":
                i = j + 1
                continue

        if line.strip() == "":
            flush_para()
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            flush_para()
            level = len(m.group(1))
            blocks.append((f"h{level}", [m.group(2).strip()]))
            i += 1
            continue

        # blockquote line
        if line.startswith(">"):
            flush_para()
            bq: list[str] = []
            while i < len(lines) and lines[i].startswith(">"):
                bq.append(lines[i][1:].lstrip())
                i += 1
            blocks.append(("quote", bq))
            continue

        # list blocks (bullets or numbers)
        if re.match(r"^\s*-\s+", line) or re.match(r"^\s*\d+\.\s+", line):
            flush_para()
            items: list[str] = []
            kind = "ul" if re.match(r"^\s*-\s+", line) else "ol"
            while i < len(lines):
                cur = lines[i]
                if kind == "ul":
                    m2 = re.match(r"^\s*-\s+(.*)$", cur)
                else:
                    m2 = re.match(r"^\s*\d+\.\s+(.*)$", cur)
                if not m2:
                    break
                items.append(m2.group(1).strip())
                i += 1
            blocks.append((kind, items))
            continue

        # markdown pipe table block
        # Expected minimal form:
        # | h1 | h2 |
        # |---|---|
        # | v1 | v2 |
        if "|" in line and i + 1 < len(lines):
            sep = lines[i + 1]
            if "|" in sep and re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", sep):
                flush_para()
                tbl: list[str] = [line.rstrip("\n"), sep.rstrip("\n")]
                i += 2
                while i < len(lines):
                    cur = lines[i]
                    if cur.strip() == "" or "|" not in cur:
                        break
                    tbl.append(cur.rstrip("\n"))
                    i += 1
                blocks.append(("table", tbl))
                continue

        para_buf.append(line.rstrip("\n"))
        i += 1

    flush_para()
    if in_code and code_buf:
        blocks.append(("code", code_buf))
    return blocks


def _add_inline_runs(paragraph, text: str, *, bold: bool = False, italic: bool = False, monospace: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if monospace:
        run.font.name = "Courier New"


def _add_markdownish_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)

    # Very small inline parser: `code`, **bold**, *italic*
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            _add_inline_runs(p, part[1:-1], monospace=True)
        elif part.startswith("**") and part.endswith("**"):
            _add_inline_runs(p, part[2:-2], bold=True)
        elif part.startswith("*") and part.endswith("*"):
            _add_inline_runs(p, part[1:-1], italic=True)
        else:
            _add_inline_runs(p, part)


def _parse_pipe_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    cells = [c.strip() for c in s.split("|")]
    return [_strip_md_inline(c) for c in cells]


def _is_sep_row(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", line))


def _add_markdown_table(doc: Document, lines: list[str]) -> None:
    if len(lines) < 2:
        return
    if not _is_sep_row(lines[1]):
        return

    header = _parse_pipe_row(lines[0])
    data_rows = [_parse_pipe_row(x) for x in lines[2:]]
    n_cols = max([len(header)] + [len(r) for r in data_rows] or [0])
    if n_cols <= 0:
        return

    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.style = "Table Grid"

    for j in range(n_cols):
        txt = header[j] if j < len(header) else ""
        p = table.cell(0, j).paragraphs[0]
        run = p.add_run(txt)
        run.bold = True

    for i, row in enumerate(data_rows, start=1):
        for j in range(n_cols):
            txt = row[j] if j < len(row) else ""
            table.cell(i, j).text = txt

    doc.add_paragraph("")


def _extract_figure_captions(md: str) -> dict[str, FigureCaption]:
    captions: dict[str, FigureCaption] = {}
    # Captions are written as: **Figure X (F1): ...** ...
    pattern = re.compile(r"^\*\*(Figure\s+\d+\s+\((F\d+)\):\s+[^*]+)\*\*\s*(.*)$")
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        m = pattern.match(line)
        if not m:
            i += 1
            continue
        title_line = _strip_md_inline(m.group(1)).strip()
        fig_id = m.group(2).strip()
        body_parts: list[str] = []
        if m.group(3).strip():
            body_parts.append(_strip_md_inline(m.group(3).strip()))
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if nxt.strip() == "":
                break
            # stop at next figure caption
            if pattern.match(nxt):
                break
            body_parts.append(_strip_md_inline(nxt))
            i += 1
        body = " ".join(x.strip() for x in body_parts if x.strip())
        captions[fig_id] = FigureCaption(figure_id=fig_id, title_line=title_line, body=body)
        i += 1
    return captions


def build_docx(md_path: Path, out_docx: Path, figures_dir: Path | None) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    blocks = _iter_blocks(md_text)
    fig_caps = _extract_figure_captions(md_text)

    doc = Document()
    # Global-ish style defaults
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title_set = False
    in_references_section = False
    for kind, content in blocks:
        if kind == "h1":
            # First h1 treated as manuscript title
            text = _strip_md_inline(content[0]).strip()
            p = doc.add_paragraph(text, style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_set = True
            continue
        if kind.startswith("h") and kind[1:].isdigit():
            level = int(kind[1:])
            text = _strip_md_inline(content[0]).strip()
            if level == 1 and not title_set:
                p = doc.add_paragraph(text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_set = True
                continue
            # Track whether we are inside the References section so that
            # numbered items there are emitted as explicit "1. ...", "2. ..."
            # text and never continue a previous Word numbered list.
            if level <= 2:
                in_references_section = text.lower() == "references"
            style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 3")
            doc.add_paragraph(text, style=style)
            continue
        if kind == "quote":
            q = "\n".join(_strip_md_inline(x) for x in content).strip()
            _add_markdownish_paragraph(doc, q, style="Quote")
            continue
        if kind == "ul":
            for item in content:
                _add_markdownish_paragraph(doc, _strip_md_inline(item), style="List Bullet")
            continue
        if kind == "ol":
            if in_references_section:
                for idx, item in enumerate(content, start=1):
                    _add_markdownish_paragraph(doc, f"{idx}. {_strip_md_inline(item)}")
            else:
                for item in content:
                    _add_markdownish_paragraph(doc, _strip_md_inline(item), style="List Number")
            continue
        if kind == "code":
            for line in content:
                p = doc.add_paragraph(style="No Spacing")
                run = p.add_run(line)
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            doc.add_paragraph("")
            continue
        if kind == "table":
            _add_markdown_table(doc, content)
            continue
        if kind == "para":
            text = " ".join(x.strip() for x in content if x.strip())
            if not text:
                continue
            # drop any residual markdown-only lines if present (defensive)
            if text.strip().startswith("<!--") and text.strip().endswith("-->"):
                continue
            _add_markdownish_paragraph(doc, text)
            continue

    # Append embedded figures at end (if available).
    # The Markdown manuscript already includes a "## Figures" section with captions;
    # we keep that content and add an explicit appendix with the rendered PNGs so
    # reviewers can see everything in one file even if the submission system
    # also requires separate figure uploads.
    if figures_dir and figures_dir.exists():
        doc.add_page_break()
        doc.add_paragraph("Appendix: Embedded Figures", style="Heading 1")
        for n in range(1, 7):
            fig_id = f"F{n}"
            png = figures_dir / f"{fig_id}_system_contract.png"
            if fig_id == "F1":
                png = figures_dir / "F1_system_contract.png"
            elif fig_id == "F2":
                png = figures_dir / "F2_reproducibility.png"
            elif fig_id == "F3":
                png = figures_dir / "F3_scrna_benchmark.png"
            elif fig_id == "F4":
                png = figures_dir / "F4_spatial_benchmark.png"
            elif fig_id == "F5":
                png = figures_dir / "F5_ops_benchmark.png"
            elif fig_id == "F6":
                png = figures_dir / "F6_robustness_matrix.png"

            if n > 1:
                doc.add_page_break()

            cap = fig_caps.get(fig_id)
            if cap:
                doc.add_paragraph(cap.title_line, style="Heading 2")
                if cap.body:
                    _add_markdownish_paragraph(doc, cap.body)
            else:
                doc.add_paragraph(f"Figure {n} ({fig_id})", style="Heading 2")

            if png.exists():
                doc.add_picture(str(png), width=Inches(6.5))
            else:
                _add_markdownish_paragraph(doc, f"[Missing figure file: {png.as_posix()}]")
            doc.add_paragraph("")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def main() -> int:
    ap = argparse.ArgumentParser(description="Export docs/MANUSCRIPT_DRAFT.md to a submission-friendly DOCX.")
    ap.add_argument("--input", default="docs/MANUSCRIPT_DRAFT.md", help="Input Markdown (default: docs/MANUSCRIPT_DRAFT.md).")
    ap.add_argument("--output", default="output/doc/CloudBioIntegrator_Manuscript.docx", help="Output DOCX path.")
    ap.add_argument(
        "--figures",
        default="plots/publication/png",
        help="Directory containing PNG figures to append at the end (default: plots/publication/png).",
    )
    args = ap.parse_args()

    md_path = (ROOT / args.input).resolve()
    out_docx = (ROOT / args.output).resolve()
    figures_dir = (ROOT / args.figures).resolve() if args.figures else None

    if not md_path.exists():
        raise SystemExit(f"missing input: {md_path}")
    build_docx(md_path, out_docx, figures_dir)
    print(f"OK: wrote {out_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
